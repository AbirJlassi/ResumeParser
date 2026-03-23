"""
CV Parser — Pipeline Hybride Stratifié
=======================================
Architecture en 4 couches ordonnées par fiabilité décroissante :

    Couche 1 — Regex         : email, téléphone, URLs (déterministe, haute précision)
    Couche 2 — NER (spaCy)   : noms propres, organisations, localisations
    Couche 3 — LLM (Groq)    : sémantique, compétences, résumé, expériences structurées
    Couche 4 — Post-processing: normalisation, validation croisée, déduplication

Formats supportés :
    - PDF (texte natif)
    - PDF image / CV scanné → OCR (Tesseract)
    - DOCX
    - TXT / Markdown
    - Image directe (JPG, PNG, WEBP, TIFF)

Dépendances :
    pip install groq pdfminer.six python-docx spacy pillow pytesseract python-dateutil
    python -m spacy download fr_core_news_sm   # ou en_core_web_sm selon la langue cible
    apt-get install tesseract-ocr tesseract-ocr-fra tesseract-ocr-eng
"""

from __future__ import annotations

import json
import logging
import os
import re
import tempfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Optional

# ── Extraction de texte ───────────────────────────────────────────────────────
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document as DocxDocument

# ── OCR ───────────────────────────────────────────────────────────────────────
from PIL import Image
import pytesseract

# ── NER ───────────────────────────────────────────────────────────────────────
import spacy

# ── LLM ───────────────────────────────────────────────────────────────────────
from groq import Groq
from dotenv import load_dotenv

# ── Normalisation des dates ───────────────────────────────────────────────────
from dateutil import parser as dateutil_parser

load_dotenv()

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("cv_parser")


# ---------------------------------------------------------------------------
# Modèle de données
# ---------------------------------------------------------------------------

@dataclass
class CVData:
    """
    Structure de sortie normalisée du parser.
    Compatible avec les systèmes ATS et les pipelines de recrutement.
    """
    full_name:      Optional[str] = None
    email:          Optional[str] = None
    phone:          Optional[str] = None
    location:       Optional[str] = None
    linkedin:       Optional[str] = None
    github:         Optional[str] = None
    summary:        Optional[str] = None
    skills:         list[str]     = field(default_factory=list)
    languages:      list[dict]    = field(default_factory=list)  # [{language, level}]
    experiences:    list[dict]    = field(default_factory=list)
    education:      list[dict]    = field(default_factory=list)
    certifications: list[str]     = field(default_factory=list)

    # Métadonnées d'extraction (non-ATS, utiles pour debugging)
    _extraction_meta: dict = field(default_factory=dict)

    def to_json(self, indent: int = 2, include_meta: bool = False) -> str:
        d = asdict(self)
        if not include_meta:
            d.pop("_extraction_meta", None)
        return json.dumps(d, ensure_ascii=False, indent=indent)

    @classmethod
    def from_dict(cls, data: dict) -> "CVData":
        return cls(
            full_name      = data.get("full_name"),
            email          = data.get("email"),
            phone          = data.get("phone"),
            location       = data.get("location"),
            linkedin       = data.get("linkedin"),
            github         = data.get("github"),
            summary        = data.get("summary"),
            skills         = data.get("skills", []),
            languages      = data.get("languages", []),
            experiences    = data.get("experiences", []),
            education      = data.get("education", []),
            certifications = data.get("certifications", []),
        )


# ---------------------------------------------------------------------------
# COUCHE 0 — Extraction de texte brut (tous formats)
# ---------------------------------------------------------------------------

class TextExtractor:
    """
    Extrait le texte brut depuis n'importe quel format de CV.

    Stratégie PDF :
        1. Tentative d'extraction native (texte sélectionnable)
        2. Si résultat insuffisant (<50 chars) → OCR via Tesseract
           (gère les PDFs scannés et les CVs en image)
    """

    # Seuil minimum de texte extrait pour considérer qu'on a du texte natif
    _PDF_MIN_CHARS = 50

    # Formats image supportés directement (OCR pur)
    _IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif", ".bmp"}

    @classmethod
    def extract(cls, file_path: str) -> tuple[str, str]:
        """
        Extrait le texte du fichier.

        Returns:
            (text, method) où method ∈ {'native', 'ocr', 'docx', 'txt'}
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return cls._extract_pdf(file_path)
        elif suffix == ".docx":
            return cls._extract_docx(file_path), "docx"
        elif suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8"), "txt"
        elif suffix in cls._IMAGE_EXTENSIONS:
            return cls._extract_image(file_path), "ocr"
        else:
            raise ValueError(f"Format non supporté : {suffix}. "
                             f"Formats acceptés : pdf, docx, txt, md, jpg, png, webp, tiff")

    @classmethod
    def _extract_pdf(cls, file_path: str) -> tuple[str, str]:
        """Extraction PDF : texte natif en priorité, OCR en fallback."""
        # Tentative 1 : extraction texte natif
        try:
            native_text = pdf_extract_text(file_path) or ""
            native_text = native_text.strip()

            if len(native_text) >= cls._PDF_MIN_CHARS:
                logger.info("PDF natif extrait (%d chars)", len(native_text))
                return native_text, "native"
            else:
                logger.info("PDF texte insuffisant (%d chars) → bascule OCR", len(native_text))
        except Exception as e:
            logger.warning("Extraction PDF native échouée : %s → bascule OCR", e)

        # Tentative 2 : OCR page par page via PIL + Tesseract
        return cls._ocr_pdf(file_path), "ocr"

    @classmethod
    def _ocr_pdf(cls, file_path: str) -> str:
        """
        OCR d'un PDF scanné.
        Convertit chaque page en image puis applique Tesseract.
        Utilise pdf2image si disponible, sinon pillow seul.
        """
        try:
            # Préférence : pdf2image (meilleure qualité)
            from pdf2image import convert_from_path  # type: ignore
            images = convert_from_path(file_path, dpi=300)
            logger.info("OCR PDF via pdf2image (%d pages)", len(images))
        except ImportError:
            # Fallback : PIL direct (qualité moindre, pas de rendu vectoriel)
            logger.warning("pdf2image non disponible — OCR dégradé. "
                           "Installer : pip install pdf2image")
            images = [Image.open(file_path)]

        pages_text = []
        for i, img in enumerate(images):
            # lang='fra+eng' couvre le français et l'anglais simultanément
            text = pytesseract.image_to_string(
                img,
                lang="fra+eng",
                config="--psm 6"   # mode : bloc de texte uniforme
            )
            pages_text.append(text)
            logger.debug("  Page %d OCR : %d chars", i + 1, len(text))

        return "\n\n".join(pages_text)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extraction DOCX : paragraphes + tableaux (souvent utilisés pour les CV)."""
        doc = DocxDocument(file_path)
        parts = []

        # Paragraphes normaux
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tableaux (certains CV structurent les infos en grille)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    @staticmethod
    def _extract_image(file_path: str) -> str:
        """OCR direct sur image (JPG, PNG, WEBP, etc.)."""
        img = Image.open(file_path)

        # Amélioration de l'image pour l'OCR
        img = img.convert("L")  # niveaux de gris
        text = pytesseract.image_to_string(
            img,
            lang="fra+eng",
            config="--psm 6"
        )
        logger.info("OCR image : %d chars extraits", len(text))
        return text


# ---------------------------------------------------------------------------
# COUCHE 1 — Regex (champs déterministes)
# ---------------------------------------------------------------------------

class RegexExtractor:
    """
    Extraction haute-précision des champs à structure syntaxique fixe.
    Toujours exécutée EN PREMIER pour verrouiller les champs critiques.

    Principe : si le pattern matche, on a une confiance maximale.
    Le résultat Regex prime toujours sur NER et LLM pour ces champs.
    """

    # ── Patterns compilés ────────────────────────────────────────────────────

    # Email : RFC 5321 simplifié (couvre 99%+ des cas réels)
    EMAIL = re.compile(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
        re.IGNORECASE
    )

    # Téléphone — deux patterns complémentaires pour couvrir tous les formats
    #
    # Bugs corrigés v2 :
    #   1. (+216) 55-987-654 : la parenthèse ouvrante était bloquée par le lookbehind
    #      → PHONE_INTL capture explicitement "(?" avant le +
    #   2. "2020-2022" matchait comme téléphone
    #      → _YEAR_RANGE invalide ces matches en post-traitement
    #   3. "+33 6 12 34 56 78" ne capturait que le corps sans indicatif
    #      → PHONE_INTL ancre sur le préfixe international en premier
    #
    # PHONE_INTL : numéros avec préfixe international (+XX, 00XX, (+XX))
    PHONE_INTL = re.compile(
        r"(?:\(?\+\d{1,4}\)?|00\d{1,3})"   # (+216), +216, 0033, +33
        r"[\s.\-]?"
        r"(?:\d[\s.\-]?)?"                    # chiffre optionnel (ex: +33 6)
        r"\d{2}(?:[\s.\-]?\d{2,3}){2,4}"    # corps du numéro
        r"(?!\d)"
    )
    # PHONE_LOCAL : numéros locaux FR/MA/TN (0X XX XX XX XX)
    PHONE_LOCAL = re.compile(
        r"(?<!\d)0[1-9](?:[\s.\-]?\d{2}){4}(?!\d)"
    )
    # Plages d'années — invalide les faux positifs (ex: "2020-2022")
    _YEAR_RANGE = re.compile(r"^(19|20)\d{2}[\s\-\u2013\u2014]+(19|20)\d{2}$")

        # URLs professionnelles
    LINKEDIN = re.compile(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/([\w\-\.]+)/?",
        re.IGNORECASE
    )
    GITHUB = re.compile(
        r"(?:https?://)?(?:www\.)?github\.com/([\w\-]+)/?",
        re.IGNORECASE
    )

    # Dates (pour validation croisée ultérieure)
    DATE_RANGE = re.compile(
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec"
        r"|jan|fév|mars|avr|mai|juin|juil|août|sep|oct|nov|déc)"
        r"[\w.]*\s+\d{4}"
        r"(?:\s*[-–—]\s*"
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec"
        r"|jan|fév|mars|avr|mai|juin|juil|août|sep|oct|nov|déc"
        r"|présent|present|aujourd|current|now)"
        r"[\w.]*(?:\s+\d{4})?)?",
        re.IGNORECASE
    )

    def extract(self, text: str) -> dict:
        """
        Extrait tous les champs déterministes du texte.

        Returns:
            dict avec clés : email, phone, linkedin, github
            None si non trouvé.
        """
        result = {}

        # Email — première occurrence (les CVs n'ont généralement qu'un email)
        email_match = self.EMAIL.search(text)
        result["email"] = email_match.group(0).lower() if email_match else None

        # Téléphone — cherche d'abord les numéros internationaux (plus spécifiques),
        # puis les numéros locaux. Retourne le premier dans l'ordre d'apparition.
        result["phone"] = self._extract_phone(text)

        # LinkedIn — on retourne l'URL normalisée
        linkedin_match = self.LINKEDIN.search(text)
        if linkedin_match:
            handle = linkedin_match.group(1)
            result["linkedin"] = f"linkedin.com/in/{handle}"
        else:
            result["linkedin"] = None

        # GitHub — idem
        github_match = self.GITHUB.search(text)
        if github_match:
            handle = github_match.group(1)
            # Exclusion des handles génériques (faux positifs fréquents)
            if handle.lower() not in ("features", "blog", "about", "contact", "login"):
                result["github"] = f"github.com/{handle}"
            else:
                result["github"] = None
        else:
            result["github"] = None

        logger.info(
            "Regex → email=%s | phone=%s | linkedin=%s | github=%s",
            bool(result["email"]),
            bool(result["phone"]),
            bool(result["linkedin"]),
            bool(result["github"]),
        )
        return result

    def _extract_phone(self, text: str) -> str | None:
        """
        Extrait le numéro de téléphone en combinant deux patterns complémentaires.

        Stratégie :
            1. PHONE_INTL  : capture les numéros avec préfixe international
               (+33, +216, (+216), 0033) — plus spécifique, traité en premier
            2. PHONE_LOCAL : capture les numéros locaux (06 XX XX XX XX)
            On retient le premier candidat valide dans l'ordre d'apparition.

        Validation :
            - Minimum 7 chiffres dans le numéro
            - Rejet des plages d'années (2020-2022) via _YEAR_RANGE
        """
        candidates = []

        for m in self.PHONE_INTL.finditer(text):
            candidate = m.group(0).strip()
            digits = re.sub(r"\D", "", candidate)
            if len(digits) >= 7 and not self._YEAR_RANGE.match(candidate):
                candidates.append((m.start(), candidate))

        for m in self.PHONE_LOCAL.finditer(text):
            candidate = m.group(0).strip()
            if not self._YEAR_RANGE.match(candidate):
                candidates.append((m.start(), candidate))

        if not candidates:
            return None

        # Retourne le premier numéro dans l'ordre d'apparition dans le texte
        candidates.sort(key=lambda x: x[0])
        return candidates[0][1]


# ---------------------------------------------------------------------------
# COUCHE 2 — NER spaCy (entités nommées)
# ---------------------------------------------------------------------------

# Tokens qui, s'ils apparaissent sur une ligne, l'excluent comme nom de candidat
_NAME_LINE_BLACKLIST: set[str] = {
    # Titres de poste (FR + EN)
    "developer", "engineer", "manager", "analyst", "architect", "consultant",
    "designer", "intern", "trainee", "officer", "lead", "senior", "junior",
    "stagiaire", "ingénieur", "développeur", "responsable", "chef", "directeur",
    # Rubriques CV
    "cv", "resume", "curriculum", "vitae", "profil", "profile", "portfolio",
    "about", "contact", "summary", "formation", "compétences", "skills",
    "experience", "expérience", "education", "projet", "projects",
    "certifications", "langues", "languages", "références", "references",
    # Technologies (souvent en majuscule initiale dans les CV)
    "python", "java", "javascript", "typescript", "react", "angular", "vue",
    "docker", "kubernetes", "aws", "azure", "gcp", "sql", "linux", "git",
    "databricks", "spark", "kafka", "hadoop", "tensorflow", "pytorch",
    # Marqueurs de contact
    "email", "phone", "tel", "mobile", "adresse", "address", "linkedin", "github",
}
 
# Mot de nom valide : lettres (dont accents), tirets, apostrophes — pas de chiffres
_NAME_WORD_RE = re.compile(r"^[A-Za-zÀ-ÿ][A-Za-z\u00C0-\u024F'\-]*$")
 
# Exclut les lignes entièrement en majuscules (rubriques : "EXPÉRIENCE PROFESSIONNELLE")
_ALL_CAPS_RE  = re.compile(r"^[A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝ\s\-]+$")
 
 
def _is_valid_name_line(line: str) -> bool:
    """
    Retourne True si la ligne est un candidat plausible pour un nom de candidat.
 
    Filtres appliqués (dans l'ordre) :
        1. Longueur entre 3 et 60 chars
        2. Pas d'email, URL, date (19xx/20xx), bullet point
        3. Pas entièrement en majuscules (= rubrique)
        4. Entre 2 et 6 mots, tous composés uniquement de lettres/tirets/apostrophes
        5. Au moins 2 mots avec initiale majuscule
        6. Aucun mot blacklisté (titres, rubriques, tech)
    """
    stripped = line.strip()
    if not stripped or len(stripped) < 3 or len(stripped) > 60:
        return False
    if re.search(r"@", stripped):                                      return False
    if re.search(r"https?://|www\.|linkedin\.com|github\.com", stripped): return False
    if re.search(r"\b(19|20)\d{2}\b", stripped):                   return False
    if re.match(r"^[\•\-\*\▸\◦\→]", stripped):                  return False
    if _ALL_CAPS_RE.match(stripped):                                    return False
 
    # Normalise les séparateurs courants (|, ·, —, /) avant de découper
    clean = re.sub(r"[|·—/]", " ", stripped)
    words = [w for w in clean.split() if w]
 
    if len(words) < 2 or len(words) > 6:
        return False
 
    # Tous les mots doivent être des "mots de nom" (pas de chiffres ni symboles)
    if not all(_NAME_WORD_RE.match(w) for w in words):
        return False
 
    # Au moins 2 mots avec initiale majuscule
    if sum(1 for w in words if w[0].isupper()) < 2:
        return False
 
    # Aucun mot blacklisté (insensible à la casse)
    if {w.lower() for w in words} & _NAME_LINE_BLACKLIST:
        return False
 
    return True
 
 
def _extract_name_from_header(text: str, max_lines: int = 10) -> str | None:
    """
    Scanne les premières lignes du CV pour extraire le nom du candidat.
 
    Heuristique :
        Le nom est presque toujours dans les 10 premières lignes non vides,
        seul sur sa ligne, sans email/URL/chiffres, avec ≥2 mots capitalisés.
        Cette approche est plus robuste que le NER spaCy pour les noms
        maghrébins, arabes ou peu représentés dans les corpus d'entraînement.
 
    Args:
        text:      Texte brut du CV
        max_lines: Nombre de lignes non vides à scanner depuis le début
 
    Returns:
        Nom du candidat ou None si non détecté
    """
    lines     = text.split("\n")
    non_empty = [l for l in lines if l.strip()][:max_lines]
    for line in non_empty:
        if _is_valid_name_line(line):
            return line.strip()
    return None
 
 
class NERExtractor:
    """
    Extraction des entités nommées via spaCy + heuristique d'en-tête.
 
    Cibles :
        - full_name → heuristique lignes d'abord, spaCy (PER) en fallback
        - location  → spaCy GPE/LOC avec fusion ville + pays contigus
        - ORG       → entreprises et établissements (pour validation LLM)
 
    Pourquoi l'heuristique en premier pour le nom :
        spaCy fr_core_news_sm est entraîné sur des dépêches AFP — il reconnaît
        mal les noms maghrébins, tunisiens, marocains. L'heuristique basée sur
        la structure du CV (nom = première ligne valide) est plus fiable en
        pratique sur des CVs francophones du Maghreb.
    """
 
    _MODEL_PRIORITY = [
        "fr_core_news_sm",
        "fr_core_news_md",
        "fr_core_news_lg",
        "en_core_web_sm",
        "en_core_web_md",
    ]
 
    _ORG_BLACKLIST = {
        "github", "linkedin", "microsoft", "google", "apple",
        "python", "javascript", "typescript", "react", "docker",
    }
 
    # Blacklist spaCy uniquement (fallback NER) — noms de tech/rubriques
    # souvent taggés PER par le modèle
    _NER_NAME_BLACKLIST = {
        "azure databricks", "google cloud", "amazon web services",
        "microsoft azure", "apache kafka", "apache spark",
        "blog personnel", "portfolio", "personal website",
        "curriculum vitae", "resume", "profil", "profile",
        "backend developer", "frontend developer", "full stack developer",
        "software engineer", "data scientist", "data engineer",
        "devops engineer", "product manager", "project manager",
        "lead developer", "senior developer", "junior developer",
    }
 
    def __init__(self):
        self.nlp = self._load_model()
 
    def _load_model(self) -> spacy.language.Language:
        for model_name in self._MODEL_PRIORITY:
            try:
                nlp = spacy.load(model_name)
                logger.info("Modèle spaCy chargé : %s", model_name)
                return nlp
            except OSError:
                continue
        logger.warning(
            "Aucun modèle spaCy installé. NER désactivé.\n"
            "  → python -m spacy download fr_core_news_sm"
        )
        return spacy.blank("fr")
 
    def extract(self, text: str) -> dict:
        """
        Extrait les entités nommées pertinentes du texte CV.
 
        Stratégie full_name (ordre de priorité) :
            1. Heuristique d'en-tête (_extract_name_from_header) — fiable pour
               les noms maghrébins et les CVs structurés classiquement
            2. spaCy NER (PER) — fallback si l'heuristique ne trouve rien
               (CVs avec mise en page atypique, formats colonnes, etc.)
 
        Returns:
            dict : full_name, location, _organizations
        """
        header_text = text[:10_000]
 
        # ── Étape 1 : heuristique d'en-tête (prioritaire) ────────────────────
        heuristic_name = _extract_name_from_header(text)
 
        # ── Étape 2 : NER spaCy ───────────────────────────────────────────────
        doc = self.nlp(header_text)
 
        ner_persons   = []
        location_ents = []
        organizations = []
 
        for ent in doc.ents:
            val   = ent.text.strip()
            label = ent.label_
 
            if not val or len(val) < 2:
                continue
 
            if label in ("PER", "PERSON"):
                words = val.split()
                if (len(words) >= 2
                        and all(w[0].isupper() for w in words if w)
                        and val.lower() not in self._NER_NAME_BLACKLIST):
                    ner_persons.append(val)
 
            elif label in ("GPE", "LOC", "LOCATION"):
                location_ents.append((ent.start_char, val))
 
            elif label in ("ORG", "ORGANIZATION"):
                if val.lower() not in self._ORG_BLACKLIST:
                    organizations.append(val)
 
        # ── Fusion des résultats ──────────────────────────────────────────────
        # Heuristique prioritaire, NER en fallback
        full_name = heuristic_name or (ner_persons[0] if ner_persons else None)
        location  = self._merge_location(location_ents, header_text)
 
        result = {
            "full_name":      full_name,
            "location":       location,
            "_organizations": list(dict.fromkeys(organizations)),
        }
 
        logger.info(
            "NER → full_name=%s (src=%s) | location=%s | orgs=%d",
            result["full_name"],
            "heuristic" if heuristic_name else "spacy",
            result["location"],
            len(result["_organizations"]),
        )
        return result
 
 
 
 
    @staticmethod
    def _merge_location(location_ents: list[tuple], source_text: str) -> str | None:
        """
        Fusionne des entités GPE consécutives en une localisation complète.
 
        Ex : "Sfax" + "Tunisie" à moins de 20 chars d'écart → "Sfax, Tunisie"
        """
        if not location_ents:
            return None
        if len(location_ents) == 1:
            return location_ents[0][1]
 
        first_pos, first_val = location_ents[0]
        search_start = first_pos + len(first_val)
        window       = source_text[search_start:search_start + 20]
        _, second_val = location_ents[1]
 
        if second_val in window:
            return f"{first_val}, {second_val}"
        return first_val

# ---------------------------------------------------------------------------
# COUCHE 3 — LLM Groq (sémantique)
# ---------------------------------------------------------------------------

class LLMExtractor:
    """
    Extraction sémantique via LLM Groq.

    Responsabilités :
        - Résumé professionnel
        - Compétences techniques (avec inférence contextuelle)
        - Expériences structurées (titre, entreprise, période, missions)
        - Formation
        - Certifications
        - Langues avec niveaux

    Le LLM reçoit en entrée :
        - Le texte brut du CV (tronqué si nécessaire)
        - Les champs déjà extraits par Regex + NER (contexte enrichi)
          → réduit les hallucinations sur les champs déjà connus
    """

    # Limites par modèle en caractères (approximation : 1 token ≈ 0.75 char FR/EN)
    #
    # Tous les modèles production Groq (mars 2026) ont 131 072 tokens de contexte.
    # On réserve ~800 tokens pour le prompt (instructions + schéma JSON),
    # soit ~97 000 chars théoriques disponibles pour le CV.
    # On plafonne à 60 000 chars en pratique : aucun CV humain réel ne dépasse ça,
    # et ça garde une marge confortable pour les futures évolutions du prompt.
    #
    # Modèles production disponibles sur Groq (source : console.groq.com/docs/models) :
    #   llama-3.1-8b-instant          560 t/s   $0.05/$0.08  — volume, rapidité
    #   llama-3.3-70b-versatile       280 t/s   $0.59/$0.79  — qualité, CVs complexes
    #   openai/gpt-oss-20b           1000 t/s   $0.075/$0.30 — ultra-rapide, haute précision
    #   openai/gpt-oss-120b           500 t/s   $0.15/$0.60  — flagship, meilleure qualité
    #
    # Modèles preview (évaluation uniquement, ne pas utiliser en production) :
    #   meta-llama/llama-4-scout-17b-16e-instruct  750 t/s  — très prometteur
    #   qwen/qwen3-32b                             400 t/s  — CVs multilingues FR/AR/EN
    #   moonshotai/kimi-k2-instruct-0905           200 t/s  — contexte 262k, CVs très longs
    #
    # Modèles dépréciés (NE PLUS UTILISER) :
    #   mixtral-8x7b-32768, llama3-8b-8192, llama3-70b-8192, gemma2-9b-it
    _CHAR_LIMITS = {
        # Production
        "llama-3.1-8b-instant":           60_000,
        "llama-3.3-70b-versatile":        60_000,
        "openai/gpt-oss-20b":             60_000,
        "openai/gpt-oss-120b":            60_000,
        # Preview
        "meta-llama/llama-4-scout-17b-16e-instruct": 60_000,
        "qwen/qwen3-32b":                 60_000,
        "moonshotai/kimi-k2-instruct-0905": 60_000,
    }
    _DEFAULT_LIMIT = 60_000  # valeur sûre pour tout nouveau modèle à 131k contexte

    # Schéma JSON que le LLM doit compléter
    _JSON_SCHEMA = """{
  "full_name":      "string | null",
  "location":       "string | null  (ville, pays)",
  "summary":        "string | null  (2-3 phrases synthétiques)",
  "skills":         ["string"],
  "languages":      [{"language": "string", "level": "string"}],
  "experiences": [
    {
      "title":       "string",
      "company":     "string",
      "location":    "string | null",
      "period":      "string  (ex: Jan 2022 - Présent)",
      "description": "string  (missions principales)"
    }
  ],
  "education": [
    {
      "degree":      "string",
      "institution": "string",
      "location":    "string | null",
      "period":      "string"
    }
  ],
  "certifications": ["string"]
}"""

    def __init__(self, model: str = "llama-3.1-8b-instant", temperature: float = 0.0):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise EnvironmentError(
                "GROQ_API_KEY manquant. Créer un fichier .env avec GROQ_API_KEY=..."
            )
        self.client    = Groq(api_key=api_key)
        self.model     = model
        self.temperature = temperature

    def extract(self, cv_text: str, pre_extracted: dict) -> dict:
        """
        Lance l'extraction LLM en fournissant le contexte pré-extrait.

        Args:
            cv_text:       Texte brut du CV
            pre_extracted: Champs déjà extraits (Regex + NER)

        Returns:
            dict structuré selon le schéma JSON
        """
        limit = self._CHAR_LIMITS.get(self.model, self._DEFAULT_LIMIT)

        # Troncature intelligente : on coupe au dernier saut de ligne
        if len(cv_text) > limit:
            truncated = cv_text[:limit].rsplit("\n", 1)[0]
            logger.warning(
                "CV tronqué : %d → %d chars (limite modèle %s)",
                len(cv_text), len(truncated), self.model
            )
        else:
            truncated = cv_text

        # Contexte pré-extrait à fournir au LLM pour réduire les hallucinations
        known = {k: v for k, v in pre_extracted.items()
                 if v and not k.startswith("_")}
        context_block = (
            f"\nINFORMATIONS DÉJÀ EXTRAITES (ne pas modifier) :\n"
            f"{json.dumps(known, ensure_ascii=False, indent=2)}\n"
            if known else ""
        )

        prompt = f"""Tu es un expert en analyse de CV et recrutement.

Analyse le CV ci-dessous et extrais les informations manquantes en respectant le schéma JSON.
{context_block}
RÈGLES :
- Réponds UNIQUEMENT avec du JSON valide, sans markdown, sans explication.
- Si une information est absente, utilise null ou [].
- Ne complète pas, n'invente pas d'informations.
- Pour les compétences, liste uniquement ce qui est explicitement mentionné.
- Normalise les dates au format "MMM YYYY" (ex: Jan 2022).

SCHÉMA JSON :
{self._JSON_SCHEMA}

CV :
---
{truncated}
---

JSON :"""

        raw = self.client.chat.completions.create(
            model    = self.model,
            messages = [{"role": "user", "content": prompt}],
            temperature = self.temperature,
        ).choices[0].message.content

        return self._parse_json_response(raw)

    @staticmethod
    def _parse_json_response(raw: str) -> dict:
        """Nettoie et parse la réponse JSON du LLM."""
        # Suppression des balises markdown
        cleaned = re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()

        # Extraction du bloc JSON
        start = cleaned.find("{")
        end   = cleaned.rfind("}") + 1
        if start == -1 or end == 0:
            logger.error("Aucun JSON dans la réponse LLM : %s", raw[:200])
            return {}

        try:
            return json.loads(cleaned[start:end])
        except json.JSONDecodeError as e:
            logger.error("JSON invalide : %s\n%s", e, cleaned[start:start+300])
            return {}


# ---------------------------------------------------------------------------
# COUCHE 4 — Post-processing & Validation croisée
# ---------------------------------------------------------------------------

class PostProcessor:
    """
    Fusion, validation croisée et normalisation des résultats des 3 couches.

    Règle de priorité par champ :
        email, phone, linkedin, github  → Regex > NER > LLM
        full_name, location             → NER > LLM  (Regex ne couvre pas)
        summary, skills, exp, edu, lang → LLM uniquement (sémantique)

    Validation croisée :
        Si le LLM a extrait un email → on vérifie qu'il apparaît dans le texte brut.
        Si non → on l'invalide et on garde le résultat Regex.
    """

    _EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

    def merge(
        self,
        regex_data: dict,
        ner_data:   dict,
        llm_data:   dict,
        raw_text:   str,
    ) -> dict:
        """
        Fusionne les résultats des 3 couches avec règles de priorité.

        Returns:
            dict final normalisé
        """
        merged = {}

        # ── Champs déterministes : Regex prime ────────────────────────────────
        for field in ("email", "phone", "linkedin", "github"):
            regex_val = regex_data.get(field)
            llm_val   = llm_data.get(field)

            if regex_val:
                # On a un match Regex fiable → on l'utilise
                merged[field] = regex_val
            elif llm_val:
                # LLM a trouvé quelque chose → validation croisée obligatoire
                merged[field] = self._validate_against_source(llm_val, raw_text, field)
            else:
                merged[field] = None

        # ── Nom : NER > LLM ───────────────────────────────────────────────────
        merged["full_name"] = (
            ner_data.get("full_name")
            or llm_data.get("full_name")
        )

        # ── Localisation : NER > LLM ──────────────────────────────────────────
        merged["location"] = (
            ner_data.get("location")
            or llm_data.get("location")
        )

        # ── Champs sémantiques : LLM uniquement ───────────────────────────────
        for field in ("summary", "skills", "languages", "experiences", "education", "certifications"):
            merged[field] = llm_data.get(field) or ([] if field != "summary" else None)

        # ── Nettoyage et normalisation ────────────────────────────────────────
        merged = self._normalize(merged)

        # ── Métadonnées de debug ──────────────────────────────────────────────
        merged["_extraction_meta"] = {
            "regex_found":  {k: bool(v) for k, v in regex_data.items()},
            "ner_found":    {k: bool(v) for k, v in ner_data.items() if not k.startswith("_")},
            "llm_overridden": self._compute_overrides(regex_data, ner_data, llm_data),
        }

        return merged

    def _validate_against_source(self, llm_value: str, raw_text: str, field: str) -> Optional[str]:
        """
        Vérifie qu'une valeur extraite par le LLM est bien présente dans le texte source.
        Évite les hallucinations sur les champs critiques.
        """
        if not llm_value:
            return None

        # Pour l'email : vérification exacte + format
        if field == "email":
            if llm_value.lower() in raw_text.lower():
                if self._EMAIL_RE.match(llm_value):
                    return llm_value
            logger.warning("Email LLM '%s' introuvable dans le texte source → invalide", llm_value)
            return None

        # Pour les autres champs : vérification de présence partielle
        if llm_value.lower() in raw_text.lower():
            return llm_value

        logger.debug("Valeur LLM '%s' non vérifiable dans la source (champ: %s)", llm_value, field)
        return llm_value  # On garde avec avertissement, pas d'invalidation stricte

    @staticmethod
    def _normalize(data: dict) -> dict:
        """Normalisation finale : types, casse, déduplication."""
        # Strings vides → None
        for f in ("full_name", "email", "phone", "location", "linkedin", "github", "summary"):
            val = data.get(f)
            if isinstance(val, str) and not val.strip():
                data[f] = None

        # Skills : lowercase, déduplication, tri
        if data.get("skills"):
            data["skills"] = sorted(set(
                s.strip().lower()
                for s in data["skills"]
                if isinstance(s, str) and s.strip()
            ))

        # Certifications : déduplication
        if data.get("certifications"):
            data["certifications"] = list(dict.fromkeys(
                c.strip() for c in data["certifications"] if isinstance(c, str) and c.strip()
            ))

        # Listes → liste vide si None
        for f in ("skills", "languages", "experiences", "education", "certifications"):
            if not isinstance(data.get(f), list):
                data[f] = []

        return data

    @staticmethod
    def _compute_overrides(regex_data: dict, ner_data: dict, llm_data: dict) -> list[str]:
        """Retourne la liste des champs où Regex/NER a écrasé le résultat LLM."""
        overrides = []
        for field in ("email", "phone", "linkedin", "github"):
            if regex_data.get(field) and llm_data.get(field):
                if regex_data[field] != llm_data.get(field):
                    overrides.append(field)
        if ner_data.get("full_name") and llm_data.get("full_name"):
            if ner_data["full_name"] != llm_data.get("full_name"):
                overrides.append("full_name")
        return overrides


# ---------------------------------------------------------------------------
# Pipeline Principal
# ---------------------------------------------------------------------------

class CVParser:
    """
    Orchestrateur du pipeline d'extraction en 4 couches.

    Usage simple :
        parser = CVParser()
        result = parser.parse("cv.pdf")
        print(result.to_json())

    Usage avec modèle spécifique :
        parser = CVParser(llm_model="llama3-70b-8192")
        result = parser.parse("cv_scanné.png")
    """

    def __init__(
        self,
        llm_model:       str   = "llama-3.1-8b-instant",
        llm_temperature: float = 0.0,
    ):
        self.text_extractor = TextExtractor()
        self.regex_extractor = RegexExtractor()
        self.ner_extractor   = NERExtractor()
        self.llm_extractor   = LLMExtractor(model=llm_model, temperature=llm_temperature)
        self.post_processor  = PostProcessor()

        logger.info("CVParser initialisé — LLM : %s", llm_model)

    def parse(self, file_path: str) -> CVData:
        """
        Parse un CV depuis un fichier.

        Args:
            file_path: Chemin vers le fichier CV (PDF, DOCX, TXT, image)

        Returns:
            CVData rempli
        """
        logger.info("═══ Parsing : %s ═══", Path(file_path).name)

        # ── Couche 0 : Extraction du texte ───────────────────────────────────
        raw_text, extraction_method = TextExtractor.extract(file_path)
        logger.info("Extraction texte : méthode=%s, longueur=%d chars",
                    extraction_method, len(raw_text))

        if len(raw_text) < 20:
            raise ValueError(
                f"Texte extrait insuffisant ({len(raw_text)} chars). "
                f"Le fichier est peut-être corrompu ou vide."
            )

        return self.parse_text(raw_text, source_path=file_path)

    def parse_text(self, raw_text: str, source_path: str = "<inline>") -> CVData:
        """
        Parse un CV depuis du texte brut déjà extrait.
        Utile pour les benchmarks et les tests unitaires.
        """
        # ── Couche 1 : Regex ──────────────────────────────────────────────────
        logger.info("Couche 1 — Regex")
        regex_data = self.regex_extractor.extract(raw_text)

        # ── Couche 2 : NER ────────────────────────────────────────────────────
        logger.info("Couche 2 — NER spaCy")
        ner_data = self.ner_extractor.extract(raw_text)

        # ── Couche 3 : LLM ────────────────────────────────────────────────────
        logger.info("Couche 3 — LLM Groq")
        pre_extracted = {**regex_data, **{k: v for k, v in ner_data.items()
                                           if not k.startswith("_")}}
        llm_data = self.llm_extractor.extract(raw_text, pre_extracted)

        # ── Couche 4 : Post-processing ────────────────────────────────────────
        logger.info("Couche 4 — Post-processing & fusion")
        merged = self.post_processor.merge(regex_data, ner_data, llm_data, raw_text)

        return CVData.from_dict(merged)


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------

class CVParserFactory:
    """
    Fabrique de parsers préconfigurés selon le cas d'usage.

    Modèles production Groq (mars 2026) — tous à 131 072 tokens de contexte :
        llama-3.1-8b-instant      560 t/s   $0.05/$0.08   volume, rapidité
        llama-3.3-70b-versatile   280 t/s   $0.59/$0.79   qualité production
        openai/gpt-oss-20b       1000 t/s   $0.075/$0.30  vitesse maximale
        openai/gpt-oss-120b       500 t/s   $0.15/$0.60   meilleure qualité

    Modèles preview (évaluation uniquement, ne pas utiliser en production) :
        meta-llama/llama-4-scout-17b-16e-instruct  750 t/s
        qwen/qwen3-32b                             400 t/s  (multilingue FR/AR/EN)
        moonshotai/kimi-k2-instruct-0905           200 t/s  (contexte 262k)

    Modèles dépréciés — NE PLUS UTILISER :
        mixtral-8x7b-32768, llama3-8b-8192, llama3-70b-8192, gemma2-9b-it
    """

    @staticmethod
    def fast() -> CVParser:
        """
        Volume et rapidité — traitement de CVs en masse.
        llama-3.1-8b-instant : 560 t/s, $0.05/1M tokens.
        Suffisant pour 95% des CVs standards (1-2 pages).
        """
        return CVParser(llm_model="llama-3.1-8b-instant")

    @staticmethod
    def accurate() -> CVParser:
        """
        Qualité en production — CVs complexes, profils seniors.
        llama-3.3-70b-versatile : meilleur ratio qualité/coût en production.
        """
        return CVParser(llm_model="llama-3.3-70b-versatile")

    @staticmethod
    def fastest() -> CVParser:
        """
        Vitesse absolue — 1000 t/s, APIs temps réel.
        openai/gpt-oss-20b : modèle le plus rapide sur Groq.
        """
        return CVParser(llm_model="openai/gpt-oss-20b")

    @staticmethod
    def best() -> CVParser:
        """
        Meilleure qualité disponible — CVs très complexes ou multilingues.
        openai/gpt-oss-120b : flagship 120B paramètres.
        """
        return CVParser(llm_model="openai/gpt-oss-120b")

    @staticmethod
    def multilingual() -> CVParser:
        """
        CVs multilingues FR/AR/EN.
        qwen/qwen3-32b — note : modèle preview, pas pour production critique.
        """
        return CVParser(llm_model="qwen/qwen3-32b")

    @staticmethod
    def custom(model: str) -> CVParser:
        """Modèle personnalisé — pour tests et expérimentations."""
        return CVParser(llm_model=model)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage : python cv_parser.py <chemin_cv> [modele]")
        print()
        print("Exemples :")
        print("  python cv_parser.py cv.pdf")
        print("  python cv_parser.py cv_scanné.jpg")
        print("  python cv_parser.py cv.docx llama3-70b-8192")
        print()
        print("Formats supportés : pdf, docx, txt, md, jpg, png, webp, tiff")
        sys.exit(1)

    path  = sys.argv[1]
    model = sys.argv[2] if len(sys.argv) > 2 else None

    parser = CVParserFactory.custom(model) if model else CVParserFactory.fast()
    result = parser.parse(path)

    print(result.to_json(include_meta=False))